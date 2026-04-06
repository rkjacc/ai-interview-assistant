"""
Resume Parser Module
Extracts text from PDF and Word documents without using LLM.
Handles PDF (.pdf) and Word (.docx) file formats.
"""

import pdfplumber
from docx import Document
import re
from pathlib import Path
from typing import Tuple


class ResumeParser:
    """Parse resume files (PDF, DOCX) and extract text content."""

    @staticmethod
    def extract_from_pdf(file_path: str) -> str:
        """
        Extract text from PDF file.
        
        Args:
            file_path: Path to PDF file
            
        Returns:
            Extracted text from PDF
        """
        text = ""
        try:
            with pdfplumber.open(file_path) as pdf:
                for page in pdf.pages:
                    text += page.extract_text() or ""
                    text += "\n"
        except Exception as e:
            raise ValueError(f"Error parsing PDF file: {str(e)}")
        
        return text.strip()

    @staticmethod
    def extract_from_docx(file_path: str) -> str:
        """
        Extract text from Word (.docx) file.
        
        Args:
            file_path: Path to DOCX file
            
        Returns:
            Extracted text from Word document
        """
        text = ""
        try:
            doc = Document(file_path)
            for paragraph in doc.paragraphs:
                text += paragraph.text + "\n"
            
            # Also extract from tables if present
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        text += cell.text + " "
                    text += "\n"
        except Exception as e:
            raise ValueError(f"Error parsing DOCX file: {str(e)}")
        
        return text.strip()

    @staticmethod
    def extract_from_file(file_path: str) -> Tuple[str, str]:
        """
        Automatically detect file type and extract text.
        
        Args:
            file_path: Path to resume file (PDF or DOCX)
            
        Returns:
            Tuple of (extracted_text, file_type)
            
        Raises:
            ValueError: If file type is not supported
        """
        file_path = Path(file_path)
        file_extension = file_path.suffix.lower()
        
        if file_extension == ".pdf":
            text = ResumeParser.extract_from_pdf(str(file_path))
            return text, "pdf"
        elif file_extension == ".docx":
            text = ResumeParser.extract_from_docx(str(file_path))
            return text, "docx"
        else:
            raise ValueError(f"Unsupported file format: {file_extension}. Supported formats: .pdf, .docx")

    @staticmethod
    def clean_text(text: str) -> str:
        """
        Clean extracted text by removing extra whitespace and normalizing formatting.
        
        Args:
            text: Raw extracted text
            
        Returns:
            Cleaned text
        """
        # Remove multiple consecutive newlines
        text = re.sub(r'\n\s*\n+', '\n', text)
        
        # Remove multiple spaces
        text = re.sub(r'  +', ' ', text)
        
        # Strip leading/trailing whitespace
        text = text.strip()
        
        return text

    @staticmethod
    def extract_sections(text: str) -> dict:
        """
        Attempt to extract resume sections (Skills, Experience, etc).
        
        Args:
            text: Cleaned resume text
            
        Returns:
            Dictionary with detected sections
        """
        sections = {
            "full_text": text,
            "skills": [],
            "experience": [],
            "education": [],
            "certifications": []
        }
        
        # Define section headers to look for
        skills_patterns = [
            r'(?:technical\s+)?skills\s*(?::|$)',
            r'(?:key\s+)?competencies\s*(?::|$)',
            r'technologies\s*(?::|$)',
        ]
        
        experience_patterns = [
            r'(?:professional\s+)?experience\s*(?::|$)',
            r'work\s+(?:history|experience)\s*(?::|$)',
            r'employment\s*(?::|$)',
        ]
        
        # Find sections and extract content
        lines = text.split('\n')
        current_section = None
        section_content = []
        
        for line in lines:
            line_lower = line.lower().strip()
            
            # Check for section headers
            if any(re.search(pattern, line_lower) for pattern in skills_patterns):
                if current_section and section_content:
                    sections[current_section].extend(section_content)
                current_section = "skills"
                section_content = []
            elif any(re.search(pattern, line_lower) for pattern in experience_patterns):
                if current_section and section_content:
                    sections[current_section].extend(section_content)
                current_section = "experience"
                section_content = []
            elif current_section and line.strip():
                section_content.append(line.strip())
        
        # Add remaining content
        if current_section and section_content:
            sections[current_section].extend(section_content)
        
        return sections
